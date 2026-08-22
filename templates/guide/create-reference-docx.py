#!/usr/bin/env python3
"""
Add custom Huawei styles to a pandoc reference DOCX.

Usage:
    # Step 1: Generate the base reference doc from pandoc
    pandoc -o guide-reference.docx --print-default-data-file reference.docx
    # Step 2: Add Huawei styles
    python3 create-reference-docx.py guide-reference.docx

Requires: python-docx (pip install python-docx)
"""

import sys
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor, Emu
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree


def add_or_get_paragraph_style(doc, name, base_style=None):
    """Get an existing paragraph style or create a new one."""
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def add_or_get_character_style(doc, name, base_style=None):
    """Get an existing character style or create a new one."""
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.CHARACTER)


def set_cell_shading(style, color_hex):
    """Add cell/paragraph shading via OXML (pPr/shd)."""
    color_hex = color_hex.replace("#", "")
    pPr = style.element.get_or_add_pPr()
    # Remove existing shd
    for existing in pPr.findall(qn("w:shd")):
        pPr.remove(existing)
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    pPr.append(shd)


def set_left_border(style, color_hex, size_pt=3):
    """Add a left paragraph border via OXML (pPr/pBdr/left)."""
    color_hex = color_hex.replace("#", "")
    size_eighth_pt = int(size_pt * 8)  # Word uses eighth-points
    pPr = style.element.get_or_add_pPr()
    # Remove existing pBdr
    for existing in pPr.findall(qn("w:pBdr")):
        pPr.remove(existing)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="{size_eighth_pt}" '
        f'w:space="4" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def set_bottom_border(style, color_hex, size_pt=1.5):
    """Add a bottom paragraph border via OXML (pPr/pBdr/bottom)."""
    color_hex = color_hex.replace("#", "")
    size_eighth_pt = int(size_pt * 8)
    pPr = style.element.get_or_add_pPr()
    for existing in pPr.findall(qn("w:pBdr")):
        pPr.remove(existing)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{size_eighth_pt}" '
        f'w:space="1" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    # Insert before spacing (OOXML order: pBdr before spacing)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is not None:
        spacing.addprevious(pBdr)
    else:
        pPr.append(pBdr)


def set_paragraph_border(style, side, size_str, color_hex):
    """Add a paragraph border via OXML (pPr/pBdr/<side>).

    side: 'bottom', 'top', 'left', 'right'
    size_str: e.g. '1.5pt' — parsed to eighth-points for w:sz
    color_hex: e.g. '000000'
    """
    color_hex = color_hex.replace("#", "")
    size_pt = float(size_str.replace("pt", ""))
    size_eighth_pt = int(size_pt * 8)
    pPr = style.element.get_or_add_pPr()
    for existing in pPr.findall(qn("w:pBdr")):
        pPr.remove(existing)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:{side} w:val="single" w:sz="{size_eighth_pt}" '
        f'w:space="1" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def set_left_indent(style, cm_value):
    """Set left indent on a paragraph style."""
    pf = style.paragraph_format
    pf.left_indent = Cm(cm_value)


def set_run_font(style, font_name, size_pt, color_hex=None, bold=False):
    """Configure font properties on a style's base run format."""
    rf = style.font
    rf.name = font_name
    rf.size = Pt(size_pt)
    if color_hex:
        rf.color.rgb = RGBColor.from_string(color_hex.replace("#", ""))
    if bold:
        rf.bold = True


def set_character_shading(style, color_hex):
    """Add run-level shading (rPr/shd) for character styles like badge."""
    color_hex = color_hex.replace("#", "")
    rPr = style.element.get_or_add_rPr()
    for existing in rPr.findall(qn("w:shd")):
        rPr.remove(existing)
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    rPr.append(shd)


def set_theme_fonts(doc, body_font):
    """Set the DOCX theme majorFont/minorFont and docDefaults to body_font.

    python-docx exposes no API for the theme part, so we reach into the
    package parts and mutate theme1.xml via lxml. This makes every style
    that references the theme (asciiTheme="minorHAnsi"/"majorHAnsi")
    inherit body_font, while styles with explicit rFonts (e.g. Source Code
    -> Cascadia Code) keep their explicit font.
    """
    A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"

    # 1. Find the theme part
    theme_part = None
    for part in doc.part.package.iter_parts():
        if str(part.partname) == "/word/theme/theme1.xml":
            theme_part = part
            break
    if theme_part is None:
        raise RuntimeError("word/theme/theme1.xml not found in package")

    # 2. Set majorFont + minorFont (latin, ea, cs) typeface
    root = etree.fromstring(theme_part.blob)
    font_scheme = root.find(f"{{{A_NS}}}themeElements/{{{A_NS}}}fontScheme")
    if font_scheme is None:
        raise RuntimeError("a:fontScheme not found in theme1.xml")
    for font_tag in ("majorFont", "minorFont"):
        group = font_scheme.find(f"{{{A_NS}}}{font_tag}")
        if group is None:
            continue
        for child_name in ("latin", "ea", "cs"):
            child = group.find(f"{{{A_NS}}}{child_name}")
            if child is None:
                child = etree.SubElement(group, f"{{{A_NS}}}{child_name}")
            child.set("typeface", body_font)
    theme_part._blob = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    # 3. Update docDefaults/rPrDefault rFonts for consistency
    dd = doc.styles.element.find(qn("w:docDefaults"))
    if dd is not None:
        rpr_default = dd.find(qn("w:rPrDefault"))
        if rpr_default is not None:
            rpr = rpr_default.find(qn("w:rPr"))
            if rpr is not None:
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is not None:
                    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                        rfonts.set(qn(f"w:{attr}"), body_font)


def fix_generated_docx(docx_path):
    """Post-process a pandoc-generated DOCX to fix heading styles.

    Pandoc overrides the reference doc's Heading styles with its own defaults
    (blue accent1 color, no border). This fixes them to match the PDF:
    near-black text, red bottom border on H1.

    Bypasses python-docx entirely — modifies styles.xml in the zip directly,
    because python-docx's save() overwrites any part blob modifications.
    """
    import zipfile, shutil
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    with zipfile.ZipFile(docx_path, 'r') as z:
        styles_xml = z.read('word/styles.xml')

    root = etree.fromstring(styles_xml)

    # Remove duplicate styleIds (keep first occurrence, remove subsequent)
    # python-docx's BabelFish lookup can create duplicates; this ensures a clean styles.xml
    seen_ids = set()
    for s in root.findall(f"{{{W_NS}}}style"):
        sid = s.get(f"{{{W_NS}}}styleId")
        if sid:
            if sid in seen_ids:
                root.remove(s)
            else:
                seen_ids.add(sid)

    for heading_id in ['Heading1', 'Heading2', 'Heading3', 'Heading4']:
        style = None
        for s in root.findall(f"{{{W_NS}}}style"):
            if s.get(f"{{{W_NS}}}styleId") == heading_id:
                style = s
                break
        if style is None:
            continue

        rPr = style.find(f"{{{W_NS}}}rPr")
        if rPr is None:
            rPr = etree.SubElement(style, f"{{{W_NS}}}rPr")

        # Fix color: Huawei red for all headings (matches PDF \color{huaweired})
        color = rPr.find(f"{{{W_NS}}}color")
        if color is not None:
            for attr in list(color.attrib.keys()):
                del color.attrib[attr]
            color.set(f"{{{W_NS}}}val", "C7000B")
        else:
            color = etree.SubElement(rPr, f"{{{W_NS}}}color")
            color.set(f"{{{W_NS}}}val", "C7000B")

        # Fix font: remove theme refs, set explicit
        rFonts = rPr.find(f"{{{W_NS}}}rFonts")
        if rFonts is not None:
            for attr in list(rFonts.attrib.keys()):
                if "Theme" in attr or "theme" in attr:
                    del rFonts.attrib[attr]
            rFonts.set(f"{{{W_NS}}}ascii", "HarmonyOS Sans")
            rFonts.set(f"{{{W_NS}}}hAnsi", "HarmonyOS Sans")

        # Add bottom border to Heading 1
        # Fix bold: H2-H4 should be regular weight (matches PDF \normalfont)
        if heading_id != 'Heading1':
            for tag in ['b', 'bCs', 'i', 'iCs']:
                for elem in rPr.findall(f"{{{W_NS}}}{tag}"):
                    rPr.remove(elem)

        # Add bottom border to Heading 1
        if heading_id == 'Heading1':
            pPr = style.find(f"{{{W_NS}}}pPr")
            if pPr is None:
                pPr = etree.SubElement(style, f"{{{W_NS}}}pPr")
            for pBdr in pPr.findall(f"{{{W_NS}}}pBdr"):
                pPr.remove(pBdr)
            pBdr = etree.Element(f"{{{W_NS}}}pBdr")
            # Insert pBdr before spacing (correct OOXML schema order: pBdr before spacing)
            spacing_elem = pPr.find(f"{{{W_NS}}}spacing")
            if spacing_elem is not None:
                spacing_elem.addprevious(pBdr)
            else:
                pPr.append(pBdr)
            bottom = etree.SubElement(pBdr, f"{{{W_NS}}}bottom")
            bottom.set(f"{{{W_NS}}}val", "single")
            bottom.set(f"{{{W_NS}}}sz", "12")  # 1.5pt = 12 eighth-points
            bottom.set(f"{{{W_NS}}}space", "1")
            bottom.set(f"{{{W_NS}}}color", "C7000B")

        # Fix size (pandoc overrides reference doc sizes — restore PDF values)
        heading_sizes = {'Heading1': '40', 'Heading2': '36', 'Heading3': '32', 'Heading4': '28'}
        target_sz = heading_sizes.get(heading_id)
        if target_sz:
            for tag in ['sz', 'szCs']:
                sz = rPr.find(f"{{{W_NS}}}{tag}")
                if sz is not None:
                    sz.set(f"{{{W_NS}}}val", target_sz)
                else:
                    sz = etree.SubElement(rPr, f"{{{W_NS}}}{tag}")
                    sz.set(f"{{{W_NS}}}val", target_sz)

        # Fix heading spacing to match PDF (guide.cls titlespacing values)
        heading_spacing = {
            'Heading1': ('0', '600'),    # before=0, after=30pt
            'Heading2': ('600', '120'),  # before=30pt, after=6pt
            'Heading3': ('200', '80'),   # before=10pt, after=4pt
            'Heading4': ('160', '80'),   # before=8pt, after=4pt
        }
        target_spacing = heading_spacing.get(heading_id)
        if target_spacing:
            pPr = style.find(f"{{{W_NS}}}pPr")
            if pPr is None:
                pPr = etree.SubElement(style, f"{{{W_NS}}}pPr")
            spacing = pPr.find(f"{{{W_NS}}}spacing")
            if spacing is None:
                spacing = etree.SubElement(pPr, f"{{{W_NS}}}spacing")
            spacing.set(f"{{{W_NS}}}before", target_spacing[0])
            spacing.set(f"{{{W_NS}}}after", target_spacing[1])

        # Add keepNext and keepLines (prevents heading from separating from next paragraph)
        pPr = style.find(f"{{{W_NS}}}pPr")
        if pPr is None:
            pPr = etree.SubElement(style, f"{{{W_NS}}}pPr")
        if pPr.find(f"{{{W_NS}}}keepNext") is None:
            keep_next = etree.Element(f"{{{W_NS}}}keepNext")
            pPr.insert(0, keep_next)
        if pPr.find(f"{{{W_NS}}}keepLines") is None:
            keep_lines = etree.Element(f"{{{W_NS}}}keepLines")
            pPr.insert(1, keep_lines)

    # Fix Heading5-9: add outlineLvl, qFormat, fix font/color to match Heading1-4 pattern
    for level in range(5, 10):
        heading_id = f"Heading{level}"
        for s in root.findall(f"{{{W_NS}}}style"):
            if s.get(f"{{{W_NS}}}styleId") == heading_id:
                pPr = s.find(f"{{{W_NS}}}pPr")
                if pPr is None:
                    pPr = etree.SubElement(s, f"{{{W_NS}}}pPr")
                if pPr.find(f"{{{W_NS}}}outlineLvl") is None:
                    outline = etree.SubElement(pPr, f"{{{W_NS}}}outlineLvl")
                    outline.set(f"{{{W_NS}}}val", str(level - 1))
                if s.find(f"{{{W_NS}}}qFormat") is None:
                    etree.SubElement(s, f"{{{W_NS}}}qFormat")
                rPr = s.find(f"{{{W_NS}}}rPr")
                if rPr is not None:
                    color = rPr.find(f"{{{W_NS}}}color")
                    if color is not None:
                        for attr in list(color.attrib.keys()):
                            del color.attrib[attr]
                        color.set(f"{{{W_NS}}}val", "1F2328")
                    rFonts = rPr.find(f"{{{W_NS}}}rFonts")
                    if rFonts is not None:
                        for attr in list(rFonts.attrib.keys()):
                            if "Theme" in attr or "theme" in attr:
                                del rFonts.attrib[attr]
                        rFonts.set(f"{{{W_NS}}}ascii", "HarmonyOS Sans")
                        rFonts.set(f"{{{W_NS}}}hAnsi", "HarmonyOS Sans")
                break

    # Fix Title style (cover page): 36pt, near-black, HarmonyOS Sans
    for s in root.findall(f"{{{W_NS}}}style"):
        if s.get(f"{{{W_NS}}}styleId") == "Title":
            rPr = s.find(f"{{{W_NS}}}rPr")
            if rPr is not None:
                color = rPr.find(f"{{{W_NS}}}color")
                if color is not None:
                    for attr in list(color.attrib.keys()):
                        del color.attrib[attr]
                    color.set(f"{{{W_NS}}}val", "1F2328")
                rFonts = rPr.find(f"{{{W_NS}}}rFonts")
                if rFonts is not None:
                    for attr in list(rFonts.attrib.keys()):
                        if "Theme" in attr or "theme" in attr:
                            del rFonts.attrib[attr]
                    rFonts.set(f"{{{W_NS}}}ascii", "HarmonyOS Sans")
                    rFonts.set(f"{{{W_NS}}}hAnsi", "HarmonyOS Sans")
                for tag in ['sz', 'szCs']:
                    sz = rPr.find(f"{{{W_NS}}}{tag}")
                    if sz is not None:
                        sz.set(f"{{{W_NS}}}val", "72")
                # Fix Title spacing for cover page (matches PDF guide.cls)
                # before=62pt (~2.2cm), after=128pt (~4.5cm)
                pPr = s.find(f"{{{W_NS}}}pPr")
                if pPr is None:
                    pPr = etree.SubElement(s, f"{{{W_NS}}}pPr")
                spacing = pPr.find(f"{{{W_NS}}}spacing")
                if spacing is None:
                    spacing = etree.SubElement(pPr, f"{{{W_NS}}}spacing")
                spacing.set(f"{{{W_NS}}}before", "1248")  # 62pt
                spacing.set(f"{{{W_NS}}}after", "2560")    # 128pt
            break

    # Fix VerbatimChar style: Consolas → Cascadia Code, 11pt → 10pt (matches PDF)
    for s in root.findall(f"{{{W_NS}}}style"):
        if s.get(f"{{{W_NS}}}styleId") == "VerbatimChar":
            rPr = s.find(f"{{{W_NS}}}rPr")
            if rPr is not None:
                rFonts = rPr.find(f"{{{W_NS}}}rFonts")
                if rFonts is not None:
                    rFonts.set(f"{{{W_NS}}}ascii", "Cascadia Code")
                    rFonts.set(f"{{{W_NS}}}hAnsi", "Cascadia Code")
                for tag in ['sz', 'szCs']:
                    sz = rPr.find(f"{{{W_NS}}}{tag}")
                    if sz is not None:
                        sz.set(f"{{{W_NS}}}val", "20")
                    else:
                        sz = etree.SubElement(rPr, f"{{{W_NS}}}{tag}")
                        sz.set(f"{{{W_NS}}}val", "20")
            break

    # Fix SourceCode style: add left/right indentation + szCs (matches PDF code block)
    for s in root.findall(f"{{{W_NS}}}style"):
        if s.get(f"{{{W_NS}}}styleId") == "SourceCode":
            pPr = s.find(f"{{{W_NS}}}pPr")
            if pPr is None:
                pPr = etree.Element(f"{{{W_NS}}}pPr")
                rPr = s.find(f"{{{W_NS}}}rPr")
                if rPr is not None:
                    rPr.addprevious(pPr)
                else:
                    s.append(pPr)
            ind = pPr.find(f"{{{W_NS}}}ind")
            if ind is None:
                ind = etree.SubElement(pPr, f"{{{W_NS}}}ind")
            ind.set(f"{{{W_NS}}}left", "397")   # 0.7cm
            ind.set(f"{{{W_NS}}}right", "284")   # 0.5cm
            # Add szCs for complex script consistency
            rPr = s.find(f"{{{W_NS}}}rPr")
            if rPr is not None:
                szCs = rPr.find(f"{{{W_NS}}}szCs")
                if szCs is None:
                    szCs = etree.SubElement(rPr, f"{{{W_NS}}}szCs")
                szCs.set(f"{{{W_NS}}}val", "20")
            break

    # Fix callout and code style spacing: 6pt before, 6pt after (matches PDF)
    for sid in ['warning', 'tip', 'infobox', 'SourceCode']:
        for s in root.findall(f"{{{W_NS}}}style"):
            if s.get(f"{{{W_NS}}}styleId") == sid:
                pPr = s.find(f"{{{W_NS}}}pPr")
                if pPr is None:
                    pPr = etree.SubElement(s, f"{{{W_NS}}}pPr")
                spacing = pPr.find(f"{{{W_NS}}}spacing")
                if spacing is None:
                    spacing = etree.SubElement(pPr, f"{{{W_NS}}}spacing")
                spacing.set(f"{{{W_NS}}}before", "120")  # 6pt
                spacing.set(f"{{{W_NS}}}after", "120")   # 6pt
                break

    # Fix justification for custom styles (python-docx doesn't persist .alignment
    # on custom styles, so we set <w:jc> directly)
    style_jc = {
        'CoverLogo': 'center',
        'CoverText': 'center',
        'CoverMeta': 'center',
        'TOCTitle': 'right',
        'ImageBlock': 'center',
    }
    for sid, jc_val in style_jc.items():
        for s in root.findall(f"{{{W_NS}}}style"):
            if s.get(f"{{{W_NS}}}styleId") == sid:
                pPr = s.find(f"{{{W_NS}}}pPr")
                if pPr is None:
                    pPr = etree.SubElement(s, f"{{{W_NS}}}pPr")
                jc = pPr.find(f"{{{W_NS}}}jc")
                if jc is None:
                    jc = etree.SubElement(pPr, f"{{{W_NS}}}jc")
                jc.set(f"{{{W_NS}}}val", jc_val)
                break

    # Fix BodyText/FirstParagraph spacing: 4pt after, 0pt before (matches PDF parskip)
    for sid in ['BodyText', 'FirstParagraph']:
        for s in root.findall(f"{{{W_NS}}}style"):
            if s.get(f"{{{W_NS}}}styleId") == sid:
                pPr = s.find(f"{{{W_NS}}}pPr")
                if pPr is not None:
                    spacing = pPr.find(f"{{{W_NS}}}spacing")
                    if spacing is not None:
                        spacing.set(f"{{{W_NS}}}after", "80")
                        spacing.set(f"{{{W_NS}}}before", "0")
                break

    # Fix all styles: add explicit font names alongside theme references
    # Prevents Word from falling back to Cambria when HarmonyOS Sans
    # is not installed (theme refs alone don't provide a fallback name)
    for style in root.findall(f"{{{W_NS}}}style"):
        rPr = style.find(f"{{{W_NS}}}rPr")
        if rPr is None:
            continue
        rFonts = rPr.find(f"{{{W_NS}}}rFonts")
        if rFonts is None:
            continue
        if rFonts.get(f"{{{W_NS}}}asciiTheme") and not rFonts.get(f"{{{W_NS}}}ascii"):
            rFonts.set(f"{{{W_NS}}}ascii", "HarmonyOS Sans")
        if rFonts.get(f"{{{W_NS}}}hAnsiTheme") and not rFonts.get(f"{{{W_NS}}}hAnsi"):
            rFonts.set(f"{{{W_NS}}}hAnsi", "HarmonyOS Sans")

    # Fix Normal style: 10.5pt (sz=21) to match PDF body text (10.5pt/14pt leading)
    for s in root.findall(f"{{{W_NS}}}style"):
        if s.get(f"{{{W_NS}}}styleId") == "Normal":
            rPr = s.find(f"{{{W_NS}}}rPr")
            if rPr is not None:
                for tag in ['sz', 'szCs']:
                    sz = rPr.find(f"{{{W_NS}}}{tag}")
                    if sz is not None:
                        sz.set(f"{{{W_NS}}}val", "21")
                    else:
                        sz = etree.SubElement(rPr, f"{{{W_NS}}}{tag}")
                        sz.set(f"{{{W_NS}}}val", "21")
            pPr = s.find(f"{{{W_NS}}}pPr")
            if pPr is not None:
                spacing = pPr.find(f"{{{W_NS}}}spacing")
                if spacing is not None:
                    spacing.set(f"{{{W_NS}}}after", "80")
                    spacing.set(f"{{{W_NS}}}line", "280")
                    spacing.set(f"{{{W_NS}}}lineRule", "atLeast")
            break

    # Fix docDefaults: 10.5pt (sz=21), spacing after=80 (4pt parskip)
    docDefaults = root.find(f"{{{W_NS}}}docDefaults")
    if docDefaults is not None:
        rPrDefault = docDefaults.find(f"{{{W_NS}}}rPrDefault")
        if rPrDefault is not None:
            rPr = rPrDefault.find(f"{{{W_NS}}}rPr")
            if rPr is not None:
                for tag in ['sz', 'szCs']:
                    sz = rPr.find(f"{{{W_NS}}}{tag}")
                    if sz is not None:
                        sz.set(f"{{{W_NS}}}val", "21")
                    else:
                        sz = etree.SubElement(rPr, f"{{{W_NS}}}{tag}")
                        sz.set(f"{{{W_NS}}}val", "21")
        pPrDefault = docDefaults.find(f"{{{W_NS}}}pPrDefault")
        if pPrDefault is not None:
            pPr = pPrDefault.find(f"{{{W_NS}}}pPr")
            if pPr is not None:
                spacing = pPr.find(f"{{{W_NS}}}spacing")
                if spacing is not None:
                    spacing.set(f"{{{W_NS}}}after", "80")
                    spacing.set(f"{{{W_NS}}}line", "280")
                    spacing.set(f"{{{W_NS}}}lineRule", "atLeast")

    # ── Add/fix Caption style for figure/table captions ─────────────────
    # PDF: \small (9pt), bold label, centered — ensure properties even if style exists
    caption_style = None
    for s in root.findall(f"{{{W_NS}}}style"):
        if s.get(f"{{{W_NS}}}styleId") == "Caption":
            caption_style = s
            break
    if caption_style is None:
        caption_style = etree.SubElement(root, f"{{{W_NS}}}style")
        caption_style.set(f"{{{W_NS}}}type", "paragraph")
        caption_style.set(f"{{{W_NS}}}styleId", "Caption")
        etree.SubElement(caption_style, f"{{{W_NS}}}name").set(f"{{{W_NS}}}val", "caption")
        etree.SubElement(caption_style, f"{{{W_NS}}}basedOn").set(f"{{{W_NS}}}val", "Normal")
        etree.SubElement(caption_style, f"{{{W_NS}}}next").set(f"{{{W_NS}}}val", "Caption")
        etree.SubElement(caption_style, f"{{{W_NS}}}uiPriority").set(f"{{{W_NS}}}val", "35")
        etree.SubElement(caption_style, f"{{{W_NS}}}qFormat")
    # Ensure pPr with spacing + centered
    pPr = caption_style.find(f"{{{W_NS}}}pPr")
    if pPr is None:
        pPr = etree.SubElement(caption_style, f"{{{W_NS}}}pPr")
    sp = pPr.find(f"{{{W_NS}}}spacing")
    if sp is None:
        sp = etree.SubElement(pPr, f"{{{W_NS}}}spacing")
    sp.set(f"{{{W_NS}}}before", "120")
    sp.set(f"{{{W_NS}}}after", "120")
    jc = pPr.find(f"{{{W_NS}}}jc")
    if jc is None:
        jc = etree.SubElement(pPr, f"{{{W_NS}}}jc")
    jc.set(f"{{{W_NS}}}val", "center")
    # Ensure rPr with bold + 9pt + HarmonyOS Sans
    rPr = caption_style.find(f"{{{W_NS}}}rPr")
    if rPr is None:
        rPr = etree.SubElement(caption_style, f"{{{W_NS}}}rPr")
    rf = rPr.find(f"{{{W_NS}}}rFonts")
    if rf is None:
        rf = etree.SubElement(rPr, f"{{{W_NS}}}rFonts")
    rf.set(f"{{{W_NS}}}ascii", "HarmonyOS Sans")
    rf.set(f"{{{W_NS}}}hAnsi", "HarmonyOS Sans")
    if rPr.find(f"{{{W_NS}}}b") is None:
        etree.SubElement(rPr, f"{{{W_NS}}}b")
    for tag in ['sz', 'szCs']:
        elem = rPr.find(f"{{{W_NS}}}{tag}")
        if elem is None:
            elem = etree.SubElement(rPr, f"{{{W_NS}}}{tag}")
        elem.set(f"{{{W_NS}}}val", "18")  # 9pt (matches PDF \small)

    # ── Add/fix badge character style (red pill, white bold text) ───────
    # PDF: bg=huaweired, white bold footnotesize (8pt)
    badge_style = None
    for s in root.findall(f"{{{W_NS}}}style"):
        if s.get(f"{{{W_NS}}}styleId") == "badge":
            badge_style = s
            break
    if badge_style is None:
        badge_style = etree.SubElement(root, f"{{{W_NS}}}style")
        badge_style.set(f"{{{W_NS}}}type", "character")
        badge_style.set(f"{{{W_NS}}}styleId", "badge")
        etree.SubElement(badge_style, f"{{{W_NS}}}name").set(f"{{{W_NS}}}val", "Badge")
        etree.SubElement(badge_style, f"{{{W_NS}}}uiPriority").set(f"{{{W_NS}}}val", "99")
    # Ensure rPr with bold + white text + red bg + 8pt + HarmonyOS Sans
    rPr = badge_style.find(f"{{{W_NS}}}rPr")
    if rPr is None:
        rPr = etree.SubElement(badge_style, f"{{{W_NS}}}rPr")
    rf = rPr.find(f"{{{W_NS}}}rFonts")
    if rf is None:
        rf = etree.SubElement(rPr, f"{{{W_NS}}}rFonts")
    rf.set(f"{{{W_NS}}}ascii", "HarmonyOS Sans")
    rf.set(f"{{{W_NS}}}hAnsi", "HarmonyOS Sans")
    if rPr.find(f"{{{W_NS}}}b") is None:
        etree.SubElement(rPr, f"{{{W_NS}}}b")
    color = rPr.find(f"{{{W_NS}}}color")
    if color is None:
        color = etree.SubElement(rPr, f"{{{W_NS}}}color")
    color.set(f"{{{W_NS}}}val", "FFFFFF")
    shd = rPr.find(f"{{{W_NS}}}shd")
    if shd is None:
        shd = etree.SubElement(rPr, f"{{{W_NS}}}shd")
    shd.set(f"{{{W_NS}}}val", "clear")
    shd.set(f"{{{W_NS}}}color", "auto")
    shd.set(f"{{{W_NS}}}fill", "C7000B")
    for tag in ['sz', 'szCs']:
        elem = rPr.find(f"{{{W_NS}}}{tag}")
        if elem is None:
            elem = etree.SubElement(rPr, f"{{{W_NS}}}{tag}")
        elem.set(f"{{{W_NS}}}val", "16")  # 8pt

    # ── Add TOC1/TOC2/TOC3 styles (Word built-in TOC entry styles) ───────
    toc_configs = [
        ("TOC1", "0",   "9"),
        ("TOC2", "420", "9"),
        ("TOC3", "840", "9"),
    ]
    for toc_id, indent, ui_pri in toc_configs:
        has_toc = any(s.get(f"{{{W_NS}}}styleId") == toc_id for s in root.findall(f"{{{W_NS}}}style"))
        if not has_toc:
            ts = etree.SubElement(root, f"{{{W_NS}}}style")
            ts.set(f"{{{W_NS}}}type", "paragraph")
            ts.set(f"{{{W_NS}}}styleId", toc_id)
            etree.SubElement(ts, f"{{{W_NS}}}name").set(f"{{{W_NS}}}val", toc_id.lower())
            etree.SubElement(ts, f"{{{W_NS}}}basedOn").set(f"{{{W_NS}}}val", "Normal")
            etree.SubElement(ts, f"{{{W_NS}}}next").set(f"{{{W_NS}}}val", "Normal")
            etree.SubElement(ts, f"{{{W_NS}}}uiPriority").set(f"{{{W_NS}}}val", ui_pri)
            etree.SubElement(ts, f"{{{W_NS}}}qFormat")
            pPr = etree.SubElement(ts, f"{{{W_NS}}}pPr")
            tabs = etree.SubElement(pPr, f"{{{W_NS}}}tabs")
            tab = etree.SubElement(tabs, f"{{{W_NS}}}tab")
            tab.set(f"{{{W_NS}}}val", "right")
            tab.set(f"{{{W_NS}}}leader", "dot")
            tab.set(f"{{{W_NS}}}pos", "9000")
            etree.SubElement(pPr, f"{{{W_NS}}}ind").set(f"{{{W_NS}}}left", indent)
            sp = etree.SubElement(pPr, f"{{{W_NS}}}spacing")
            sp.set(f"{{{W_NS}}}after", "40")  # 2pt

    modified_xml = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    # Replace styles.xml + numbering.xml in the DOCX zip
    # (settings.xml left untouched — updateFields triggers a Word security prompt)

    # ── Fix list indentation in numbering.xml (match PDF 1.6em/1.8em) ────
    modified_numbering = None
    with zipfile.ZipFile(docx_path, 'r') as z:
        if 'word/numbering.xml' in z.namelist():
            numbering_xml = z.read('word/numbering.xml')
            num_root = etree.fromstring(numbering_xml)
            # PDF: itemize leftmargin=1.6em ≈ 336tw, enumerate=1.8em ≈ 378tw
            # Word: left=text position, hanging=distance to bullet
            # Use left=480/hanging=240 for level 0 (text at 0.85cm, bullet at 0.42cm)
            indent_configs = [
                (0, "480", "240"),   # level 0: text=480tw, bullet=240tw
                (1, "960", "240"),   # level 1: text=960tw, bullet=720tw
                (2, "1440", "240"),  # level 2: text=1440tw, bullet=1200tw
            ]
            for lvl_num, left_val, hang_val in indent_configs:
                for lvl in num_root.iter(f"{{{W_NS}}}lvl"):
                    if lvl.get(f"{{{W_NS}}}ilvl") == str(lvl_num):
                        # w:ind is inside w:pPr inside w:lvl
                        ind = lvl.find(f"{{{W_NS}}}pPr/{{{W_NS}}}ind")
                        if ind is not None:
                            ind.set(f"{{{W_NS}}}left", left_val)
                            ind.set(f"{{{W_NS}}}hanging", hang_val)
            modified_numbering = etree.tostring(
                num_root, xml_declaration=True, encoding="UTF-8", standalone=True
            )

    tmp_path = docx_path + '.tmp'
    with zipfile.ZipFile(docx_path, 'r') as zin:
        with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == 'word/styles.xml':
                    zout.writestr(item, modified_xml)
                elif item.filename == 'word/numbering.xml' and modified_numbering is not None:
                    zout.writestr(item, modified_numbering)
                else:
                    zout.writestr(item, zin.read(item.filename))
    shutil.move(tmp_path, docx_path)
    print(f"✓ Fixed heading styles in {docx_path}")


def main():
    if len(sys.argv) >= 2 and sys.argv[1] == "--fix":
        if len(sys.argv) < 3:
            print(f"Usage: {sys.argv[0]} --fix <generated.docx>")
            sys.exit(1)
        fix_generated_docx(sys.argv[2])
        return
    if len(sys.argv) < 2:
        print(f"Usage: {sys.argv[0]} <reference.docx>")
        sys.exit(1)

    docx_path = sys.argv[1]
    doc = Document(docx_path)

    # ── Warning callout ──────────────────────────────────────────────
    style = add_or_get_paragraph_style(doc, "warning")
    set_left_border(style, "F57C00", size_pt=3)
    set_cell_shading(style, "FFF8E1")
    set_left_indent(style, 0.5)

    # ── Tip callout ──────────────────────────────────────────────────
    style = add_or_get_paragraph_style(doc, "tip")
    set_left_border(style, "2E7D32", size_pt=3)
    set_cell_shading(style, "E8F5E9")
    set_left_indent(style, 0.5)

    # ── Info callout ─────────────────────────────────────────────────
    style = add_or_get_paragraph_style(doc, "infobox")
    set_left_border(style, "1565C0", size_pt=3)
    set_cell_shading(style, "E3F2FD")
    set_left_indent(style, 0.5)

    # ── Objectives block ─────────────────────────────────────────────
    add_or_get_paragraph_style(doc, "objectives")

    # ── Changelog section ────────────────────────────────────────────
    add_or_get_paragraph_style(doc, "changelog")

    # ── Huawei table ─────────────────────────────────────────────────
    add_or_get_paragraph_style(doc, "hutable")

    # ── Source Code ──────────────────────────────────────────────────
    style = add_or_get_paragraph_style(doc, "Source Code")
    set_cell_shading(style, "F6F8FA")
    set_run_font(style, "Cascadia Code", 10, color_hex="1F2328")

    # ── Badge (character style) ──────────────────────────────────────
    style = add_or_get_character_style(doc, "badge")
    set_character_shading(style, "C7000B")
    set_run_font(style, "HarmonyOS Sans", 9, color_hex="FFFFFF", bold=True)

    # ── CoverLogo (paragraph style) ───────────────────────────────────
    style = add_or_get_paragraph_style(doc, "CoverLogo")
    style.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = style.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(10)

    # ── CoverText (paragraph style) ──────────────────────────────────
    style = add_or_get_paragraph_style(doc, "CoverText")
    style.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = style.paragraph_format
    pf.space_before = Pt(30)
    pf.space_after = Pt(10)
    set_run_font(style, "HarmonyOS Sans", 16, color_hex="1F2328")

    # ── CoverMeta (paragraph style) ──────────────────────────────────
    style = add_or_get_paragraph_style(doc, "CoverMeta")
    style.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = style.paragraph_format
    pf.space_before = Pt(5)
    pf.space_after = Pt(10)
    set_run_font(style, "HarmonyOS Sans", 12, color_hex="595959")

    # ── TOCTitle (paragraph style) ───────────────────────────────────
    style = add_or_get_paragraph_style(doc, "TOCTitle")
    style.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(10)
    set_run_font(style, "HarmonyOS Sans", 22, color_hex="1F2328", bold=True)

    # ── ImageBlock (paragraph style) ─────────────────────────────────
    style = add_or_get_paragraph_style(doc, "ImageBlock")
    style.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = style.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)

    # ── ObjectivesRule (paragraph style) ─────────────────────────────
    style = add_or_get_paragraph_style(doc, "ObjectivesRule")
    pf = style.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(10)
    set_paragraph_border(style, "bottom", "1.5pt", "000000")

    # NOTE: Heading 1-4 styles are NOT modified here. python-docx's BabelFish
    # lookup fails on pandoc's reference DOCX (case sensitivity), creating
    # duplicate styleIds. All heading style fixes are handled by
    # fix_generated_docx() which modifies styles.xml directly.

    # ── Title (cover page) — 36pt, near-black, bold, centered (matches PDF) ──
    try:
        title_style = doc.styles["Title"]
    except KeyError:
        title_style = doc.styles.add_style("Title", WD_STYLE_TYPE.PARAGRAPH)
    set_run_font(title_style, "HarmonyOS Sans", 36, color_hex="1F2328", bold=True)

    # ── Verbatim Char — Cascadia Code (matches PDF code font) ──────────────
    try:
        vc_style = doc.styles["Verbatim Char"]
    except KeyError:
        vc_style = doc.styles.add_style("Verbatim Char", WD_STYLE_TYPE.CHARACTER)
    set_run_font(vc_style, "Cascadia Code", 11, color_hex="1F2328")

    # ── Hyperlink ────────────────────────────────────────────────────
    try:
        hl = doc.styles["Hyperlink"]
    except KeyError:
        hl = doc.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
    set_run_font(hl, "HarmonyOS Sans", 10.5, color_hex="0000FF")
    # Remove underline via OXML
    rPr = hl.element.get_or_add_rPr()
    for existing in rPr.findall(qn("w:u")):
        rPr.remove(existing)
    u_elem = parse_xml(f'<w:u {nsdecls("w")} w:val="none"/>')
    rPr.append(u_elem)

    # ── Theme + default fonts: HarmonyOS Sans for all body/heading text ──
    set_theme_fonts(doc, "HarmonyOS Sans")

    # ── Page layout: A4, margins 3/3/2/2 cm (matches PDF \geometry) ──────
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Different first page: cover page has no header/footer
    section.different_first_page_header_footer = True

    # ── Header: document title via STYLEREF field (10pt, centered) ────────
    # STYLEREF "Title" automatically shows the text of the first paragraph
    # with the Title style — no need to know the title at build time.
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in list(hp.runs):
        run._element.getparent().remove(run._element)
    for run in [hp.add_run(), hp.add_run(), hp.add_run(), hp.add_run("Document Title"), hp.add_run()]:
        run.font.size = Pt(10)
        run.font.name = "HarmonyOS Sans"
    # Field: STYLEREF "Title"
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    hp.runs[0]._r.append(fld_begin)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' STYLEREF "Title" \\* MERGEFORMAT '
    hp.runs[1]._r.append(instr)
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    hp.runs[2]._r.append(fld_sep)
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    hp.runs[4]._r.append(fld_end)

    # ── Footer: page number (10pt, centered) ──────────────────────────────
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in list(fp.runs):
        run._element.getparent().remove(run._element)
    for run in [fp.add_run(), fp.add_run(), fp.add_run()]:
        run.font.size = Pt(10)
        run.font.name = "HarmonyOS Sans"
    # Field: PAGE
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    fp.runs[0]._r.append(fld_begin)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fp.runs[1]._r.append(instr)
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    fp.runs[2]._r.append(fld_end)

    # ── Save ─────────────────────────────────────────────────────────
    doc.save(docx_path)
    print(f"✓ Huawei styles added to {docx_path}")


if __name__ == "__main__":
    main()
