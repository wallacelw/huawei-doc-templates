#!/usr/bin/env python3
"""Reusable Huawei Cloud technical-report library (DOCX engine).

Provides constants, helpers, and content builders for creating Huawei-branded
technical reports as .docx files using python-docx. Built on the technical
report template.

Usage:
    from huawei_technical import new_report, add_heading, add_table, add_callout, ...

Brand colors are locked (AGENTS.md L9). Callout names are locked (L3):
warning, tip, infobox.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import os, subprocess, copy

# Default template path (auto-detected relative to this file)
_DEFAULT_TEMPLATE = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    'common-assets', 'technical-report-template.docx'
)

# ── Brand colors (locked — AGENTS.md L9) ────────────────────────────
HUAWEI_RED = RGBColor(0xC7, 0x00, 0x0B)
CODE_BG = RGBColor(0xF6, 0xF8, 0xFA)
CODE_TEXT = RGBColor(0x1F, 0x23, 0x28)
LINK_BLUE = RGBColor(0x00, 0x00, 0xFF)
RULE_BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x1F, 0x23, 0x28)
GRAY_BG = RGBColor(0xF6, 0xF8, 0xFA)

# ── Callout colors (Material Design — AGENTS.md L9, matches huawei-colors.sty)
# Warning (amber)
WARNING_BG = RGBColor(0xFF, 0xF8, 0xE1)
WARNING_FG = RGBColor(0xF5, 0x7C, 0x00)
WARNING_BD = RGBColor(0xF5, 0x7C, 0x00)
# Tip (green)
TIP_BG = RGBColor(0xE8, 0xF5, 0xE9)
TIP_FG = RGBColor(0x2E, 0x7D, 0x32)
TIP_BD = RGBColor(0x2E, 0x7D, 0x32)
# Infobox (blue)
INFO_BG = RGBColor(0xE3, 0xF2, 0xFD)
INFO_FG = RGBColor(0x15, 0x65, 0xC0)
INFO_BD = RGBColor(0x15, 0x65, 0xC0)

# Convenience aliases matching the PPT library naming
AMBER_BG = WARNING_BG
AMBER_FG = WARNING_FG
AMBER_BD = WARNING_BD
GREEN_BG = TIP_BG
GREEN_FG = TIP_FG
GREEN_BD = TIP_BD
BLUE_BG = INFO_BG
BLUE_FG = INFO_FG
BLUE_BD = INFO_BD


# ── Template loading ────────────────────────────────────────────────

def load_template(template_path=None):
    """Load the technical report template and return a Document object.

    Args:
        template_path: Path to the .docx template. Defaults to the bundled
            ``common-assets/technical-report-template.docx``.

    Returns:
        A python-docx Document object loaded from the template.
    """
    if template_path is None:
        template_path = _DEFAULT_TEMPLATE
    return Document(template_path)


def new_report(template_path=None):
    """Create a new document from the technical report template.

    This is the primary entry point for creating a Huawei-branded technical
    report. The returned document inherits all styles, sections, and
    formatting from the bundled template.

    Args:
        template_path: Path to the .docx template. Defaults to the bundled
            ``common-assets/technical-report-template.docx``.

    Returns:
        A python-docx Document object ready for content.
    """
    return load_template(template_path)


# ── Heading ─────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    """Add a heading with Huawei red color for level 1.

    Level 1 headings get Huawei red color applied to all runs. Other levels
    use the template's default heading style.

    Args:
        doc: Document object.
        text: Heading text.
        level: Heading level (1–4). Default 1.

    Returns:
        The heading paragraph.
    """
    heading = doc.add_heading(text, level=level)
    if level == 1:
        for run in heading.runs:
            run.font.color.rgb = HUAWEI_RED
    return heading


# ── Paragraph ───────────────────────────────────────────────────────

def add_paragraph(doc, text, style=None):
    """Add a paragraph to the document.

    Args:
        doc: Document object.
        text: Paragraph text.
        style: Paragraph style name or object. Defaults to 'Normal'.

    Returns:
        The paragraph object.
    """
    if style is None:
        style = 'Normal'
    return doc.add_paragraph(text, style=style)


# ── Table ───────────────────────────────────────────────────────────

def _set_cell_shading(cell, color_hex):
    """Set the background shading of a table cell using XML.

    Args:
        cell: Table cell object.
        color_hex: Hex color string without '#' (e.g. 'C7000B').
    """
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _rgb_to_hex(color):
    """Convert an RGBColor to a hex string without '#'.

    Args:
        color: RGBColor object.

    Returns:
        Hex string like 'C7000B'.
    """
    return f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'


def add_table(doc, headers, rows):
    """Add a table with Huawei red header, alternating row colors, first column bold.

    Matches the hutable style from the LaTeX template: Huawei-red header bar
    with white bold text, alternating white / light-gray body rows, and bold
    first column in data rows.

    Args:
        doc: Document object.
        headers: List of header cell strings.
        rows: List of row lists (each row is a list of cell strings).

    Returns:
        The table object.
    """
    nc = len(headers)
    nr = len(rows) + 1  # +1 for header
    table = doc.add_table(rows=nr, cols=nc)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table borders (full grid in Huawei red)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="C7000B"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="C7000B"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="C7000B"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="C7000B"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="C7000B"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="C7000B"/>'
        f'</w:tblBorders>'
    )
    # Remove existing borders if any, then add new ones
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)

    # Style header row
    for ci, h in enumerate(headers):
        cell = table.cell(0, ci)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(cell, _rgb_to_hex(HUAWEI_RED))

    # Style data rows
    for ri, row in enumerate(rows):
        bg = GRAY_BG if ri % 2 == 0 else WHITE
        bg_hex = _rgb_to_hex(bg)
        for ci, val in enumerate(row):
            cell = table.cell(ri + 1, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.color.rgb = DARK
            run.font.size = Pt(10)
            if ci == 0:
                run.font.bold = True
            _set_cell_shading(cell, bg_hex)

    return table


# ── Callout boxes ───────────────────────────────────────────────────

def add_callout(doc, kind, text):
    """Add a callout box as a single-cell table with colored background.

    Callout names are locked (AGENTS.md L3): ``warning``, ``tip``,
    ``infobox``.

    The callout is rendered as a single-cell table with a colored background,
    a 3pt left border in the accent color, and a bold label prefix.

    Args:
        doc: Document object.
        kind: One of ``'warning'``, ``'tip'``, ``'infobox'``.
        text: Callout text content.

    Returns:
        The table object representing the callout box.
    """
    colors = {
        'warning': (WARNING_BG, WARNING_FG, WARNING_BD, 'Important'),
        'tip':     (TIP_BG,     TIP_FG,     TIP_BD,     'Tip'),
        'infobox': (INFO_BG,    INFO_FG,    INFO_BD,    'Info'),
    }
    bg, fg, bd, label = colors.get(kind, colors['infobox'])

    # Create a 1x1 table
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)

    # Set cell shading (background color)
    _set_cell_shading(cell, _rgb_to_hex(bg))

    # Set left border (3pt accent color), thin other borders
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="2" w:space="0" w:color="{_rgb_to_hex(bg)}"/>'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{_rgb_to_hex(bd)}"/>'
        f'  <w:bottom w:val="single" w:sz="2" w:space="0" w:color="{_rgb_to_hex(bg)}"/>'
        f'  <w:right w:val="single" w:sz="2" w:space="0" w:color="{_rgb_to_hex(bg)}"/>'
        f'</w:tcBorders>'
    )
    # Remove existing borders if any
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)

    # Set cell margins
    cell_margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="60" w:type="dxa"/>'
        f'  <w:left w:w="120" w:type="dxa"/>'
        f'  <w:bottom w:w="60" w:type="dxa"/>'
        f'  <w:right w:w="120" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    existing_mar = tcPr.find(qn('w:tcMar'))
    if existing_mar is not None:
        tcPr.remove(existing_mar)
    tcPr.append(cell_margins)

    # Add label + text
    cell.text = ''
    p = cell.paragraphs[0]

    # Label run (bold, colored)
    label_run = p.add_run(f'{label}: ')
    label_run.font.bold = True
    label_run.font.color.rgb = fg
    label_run.font.size = Pt(10)

    # Text run
    text_run = p.add_run(text)
    text_run.font.color.rgb = fg
    text_run.font.size = Pt(10)

    # Remove table borders (the cell borders handle the visual)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblPr.append(borders)

    return table


# ── Placeholder replacement ─────────────────────────────────────────

import re

# Pattern to match {{PLACEHOLDER}} markers
_PLACEHOLDER_RE = re.compile(r'\{\{(\w+)\}\}')


def fill_section(doc, placeholder, text):
    """Replace a placeholder text in the template with actual content.

    Searches all paragraphs in the document for the exact placeholder string
    and replaces it with the given text. If the placeholder appears in
    multiple paragraphs, all occurrences are replaced.

    Args:
        doc: Document object.
        placeholder: The placeholder string to find (e.g. '[Product Name]').
        text: The replacement text.

    Returns:
        The number of replacements made.
    """
    count = 0
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, text)
                    count += 1
            # Also check if placeholder spans multiple runs
            if placeholder in paragraph.text and count == 0:
                # Rebuild from full text
                full_text = paragraph.text
                if placeholder in full_text:
                    new_text = full_text.replace(placeholder, text)
                    # Clear all runs and set text in first run
                    for run in paragraph.runs:
                        run.text = ''
                    if paragraph.runs:
                        paragraph.runs[0].text = new_text
                    else:
                        paragraph.add_run(new_text)
                    count += 1
    return count


def fill_template(doc, replacements):
    """Fill placeholder markers in a template document.

    Placeholders in the document use the ``{{NAME}}`` format. This function
    finds all occurrences in paragraphs and table cells, and replaces them
    with the corresponding values from the replacements dict.

    For multi-line values (containing ``\\n``), the first line replaces the
    placeholder in the current paragraph, and subsequent lines are inserted
    as new paragraphs immediately after it.

    Args:
        doc: python-docx Document object.
        replacements: dict mapping placeholder names (without the ``{{}}``
            delimiters) to text values, e.g.::

                {
                    'PROBLEM_DESCRIPTION': 'The customer reports...',
                    'ROOT_CAUSE': 'The issue is caused by...',
                    'WORKAROUND': '1. Log in to...\\n2. Navigate to...',
                }

    Returns:
        The total number of placeholders replaced.
    """
    count = 0

    # Process paragraphs (iterate in reverse so insertions don't shift indices)
    paragraphs = doc.paragraphs
    for i in range(len(paragraphs) - 1, -1, -1):
        para = paragraphs[i]
        matches = list(_PLACEHOLDER_RE.finditer(para.text))
        if not matches:
            continue

        # Collect all placeholder names in this paragraph
        for match in matches:
            name = match.group(1)
            if name not in replacements:
                continue

            replacement = replacements[name]
            placeholder_str = match.group(0)

            # Handle multi-line: split on \n
            lines = replacement.split('\n')
            first_line = lines[0]

            # Replace in runs
            _replace_in_paragraph(para, placeholder_str, first_line)
            count += 1

            # Insert additional lines as new paragraphs after current one
            if len(lines) > 1:
                # Get the XML element of the current paragraph
                current_elem = para._element
                parent = current_elem.getparent()
                insert_after = current_elem
                for extra_line in lines[1:]:
                    # Clone the paragraph element for consistent formatting
                    new_para_elem = copy.deepcopy(current_elem)
                    # Clear its text and set the new line
                    for r_elem in new_para_elem.findall(qn('w:r')):
                        new_para_elem.remove(r_elem)
                    # Create a new run with the text
                    # Copy run properties from original first run if available
                    r_elem = parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{_xml_escape(extra_line)}</w:t></w:r>')
                    # Try to copy rPr from original paragraph's first run
                    orig_runs = current_elem.findall(qn('w:r'))
                    if orig_runs:
                        orig_rPr = orig_runs[0].find(qn('w:rPr'))
                        if orig_rPr is not None:
                            r_elem.insert(0, copy.deepcopy(orig_rPr))
                    new_para_elem.append(r_elem)
                    # Insert after the current position
                    insert_after.addnext(new_para_elem)
                    insert_after = new_para_elem

    # Process table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    matches = list(_PLACEHOLDER_RE.finditer(para.text))
                    if not matches:
                        continue
                    for match in matches:
                        name = match.group(1)
                        if name not in replacements:
                            continue
                        replacement = replacements[name]
                        placeholder_str = match.group(0)
                        # For table cells, replace newlines with spaces
                        # (multi-line in cells is problematic)
                        cell_text = replacement.replace('\n', ' ')
                        _replace_in_paragraph(para, placeholder_str, cell_text)
                        count += 1

    return count


def _replace_in_paragraph(para, placeholder, text):
    """Replace a placeholder string in a paragraph's runs.

    Handles the case where the placeholder may span multiple runs by
    consolidating runs when needed.

    Args:
        para: Paragraph object.
        placeholder: The placeholder string (e.g. '{{PROBLEM_DESCRIPTION}}').
        text: The replacement text.
    """
    # First try simple per-run replacement
    for run in para.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, text)
            return

    # Placeholder spans multiple runs — consolidate
    if placeholder in para.text:
        full_text = para.text.replace(placeholder, text)
        # Clear all runs and set text in first run
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = full_text
        else:
            para.add_run(full_text)


def _xml_escape(text):
    """Escape special XML characters in text.

    Args:
        text: Plain text string.

    Returns:
        XML-safe string.
    """
    return (text
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;')
            .replace("'", '&apos;'))


def create_technical_report(replacements, template_path=None):
    """Create a complete technical report from the template.

    Loads the technical report template, fills all placeholder markers with
    the provided values, and returns the document ready for saving.

    Args:
        replacements: dict of placeholder name → text values. Supported
            placeholders::

                TITLE               — Report title on the cover page
                RELEASE_DATE        — Release date on the cover page
                PROBLEM_DESCRIPTION  — Problem Description and Impact section
                ROOT_CAUSE_ANALYSIS — Root Cause Analysis section
                ROOT_CAUSE          — Root Cause section
                TRIGGER_CONDITION   — Trigger Condition section
                IMPACT              — 5.1 Impact subsection
                BACKUP_DATA         — 5.2 Back up data subsection
                WORKAROUND          — 5.3 Workaround subsection
                VERIFICATION        — 5.4 Verification subsection
                ROLLBACK            — 5.5 Rollback Operation subsection
                CLEANUP             — 5.6 Cleanup Operation subsection
                VERSION             — Version info table cell
                SCENARIO            — Installation Scenario table cell

        template_path: optional custom template path. Defaults to the
            bundled ``technical-report-template.docx``.

    Returns:
        Document object ready to save with :func:`save_report`.
    """
    doc = load_template(template_path)
    fill_template(doc, replacements)
    return doc


# Backward-compatibility alias (the template was renamed from "analysis" to
# "technical"). Existing generators that call create_analysis_report() keep
# working without changes.
create_analysis_report = create_technical_report


# ── Save / export ───────────────────────────────────────────────────

def save_report(doc, filepath):
    """Save the document to a file.

    Args:
        doc: Document object.
        filepath: Output path for the .docx file.

    Returns:
        The absolute path to the saved file.
    """
    filepath = os.path.abspath(filepath)
    doc.save(filepath)
    return filepath


def to_pdf(docx_path):
    """Convert a .docx file to PDF using LibreOffice.

    Args:
        docx_path: Path to the .docx file.

    Returns:
        The path to the generated .pdf file (same directory, same stem).

    Raises:
        FileNotFoundError: If LibreOffice is not installed.
        subprocess.CalledProcessError: If conversion fails.
    """
    docx_path = os.path.abspath(docx_path)
    outdir = os.path.dirname(docx_path)
    result = subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'pdf',
         '--outdir', outdir, docx_path],
        capture_output=True, text=True, check=True
    )
    pdf_path = os.path.splitext(docx_path)[0] + '.pdf'
    return pdf_path
